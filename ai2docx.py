#!/usr/bin/env python3
"""
AI2DOCX — 将 AI 生成内容（含 LaTeX 数学公式、Markdown 排版）一键转为 Word 文档

核心功能：
  - ✅ AI 内容（DeepSeek/ChatGPT/Kimi/豆包等）→ Word
  - ✅ LaTeX 公式 → Word 原生可编辑公式
  - ✅ Markdown 排版（标题/列表/代码/表格）
  - ✅ 代码块语法高亮样式
  - ✅ 导出 .docx
  - ✅ 纯本地处理，数据不上传

使用方式：
  python ai2docx.py                       # 从剪贴板读取 → 输出 output.docx
  python ai2docx.py input.md              # 从文件读取
  python ai2docx.py input.md output.docx  # 指定输出路径
  python ai2docx.py --out result.docx     # 从剪贴板 + 指定输出

依赖（已安装）：
  - pandoc     : Markdown+LaTeX → Word 的核心转换引擎
  - python-docx: 文档后处理（样式微调、Mermaid 标注等）
  - pyperclip  : 剪贴板读取
"""

import argparse
import re
import subprocess
import sys
import tempfile
import shutil
import os
import html
from pathlib import Path

# 修复 Windows 终端编码问题（避免 GBK 无法显示 emoji 报错）
if sys.platform == "win32":
    sys.stdin.reconfigure(encoding="utf-8", errors="replace")
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# ============================================================
# Windows HTML 剪贴板提取（表格结构化来源）
# ============================================================

_DEBUG = False  # 由 --debug 标志控制


def _get_html_from_clipboard() -> str | None:
    """
    从 Windows 剪贴板提取 HTML 格式内容。
    AI 网页复制时 HTML 格式保留了完整的表格结构，而纯文本可能丢失。
    """
    if sys.platform != "win32":
        return None
    try:
        import ctypes
        from ctypes import wintypes

        kernel32 = ctypes.windll.kernel32
        user32 = ctypes.windll.user32

        # 64 位兼容：显式设置所有参数和返回类型
        user32.RegisterClipboardFormatW.restype = wintypes.UINT
        user32.OpenClipboard.restype = wintypes.BOOL
        user32.OpenClipboard.argtypes = [wintypes.HWND]
        user32.GetClipboardData.restype = wintypes.HANDLE
        user32.GetClipboardData.argtypes = [wintypes.UINT]
        user32.CloseClipboard.restype = wintypes.BOOL
        kernel32.GlobalLock.restype = wintypes.LPVOID
        kernel32.GlobalLock.argtypes = [wintypes.HGLOBAL]
        kernel32.GlobalSize.restype = ctypes.c_size_t
        kernel32.GlobalSize.argtypes = [wintypes.HGLOBAL]
        kernel32.GlobalUnlock.restype = wintypes.BOOL
        kernel32.GlobalUnlock.argtypes = [wintypes.HGLOBAL]

        CF_HTML = user32.RegisterClipboardFormatW("HTML Format")
        if CF_HTML == 0:
            return None

        if not user32.OpenClipboard(None):
            return None

        try:
            hdata = user32.GetClipboardData(CF_HTML)
            if not hdata:
                return None

            size = kernel32.GlobalSize(hdata)
            if size == 0:
                return None

            ptr = kernel32.GlobalLock(hdata)
            if not ptr:
                return None

            try:
                buf = ctypes.create_string_buffer(size)
                ctypes.memmove(buf, ptr, size)
                raw = buf.raw
                # 找到 null 终止符
                null_pos = raw.find(b"\x00")
                if null_pos >= 0:
                    raw = raw[:null_pos]
                try:
                    return raw.decode("utf-8")
                except UnicodeDecodeError:
                    return raw.decode("gbk", errors="replace")
            finally:
                kernel32.GlobalUnlock(hdata)
        finally:
            user32.CloseClipboard()
    except Exception:
        return None


def _extract_tables_from_html(html_content: str) -> list[list[list[str]]] | None:
    """
    从 HTML 中提取所有表格，返回二维字符串数组的列表。
    每个表格为 rows × cols 的字符串矩阵。
    """
    from html.parser import HTMLParser

    class TableExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.tables: list[list[list[str]]] = []
            self._current_table: list[list[str]] | None = None
            self._current_row: list[str] | None = None
            self._current_cell: str | None = None
            self._in_cell = False
            self._skip_tags = {"script", "style", "head", "meta", "link"}

        def handle_starttag(self, tag, attrs):
            tagl = tag.lower()
            if tagl in self._skip_tags:
                return
            if tagl == "table":
                self._current_table = []
            elif tagl == "tr":
                if self._current_table is not None:
                    self._current_row = []
            elif tagl in ("td", "th"):
                if self._current_row is not None:
                    self._in_cell = True
                    self._current_cell = ""

        def handle_endtag(self, tag):
            tagl = tag.lower()
            if tagl in self._skip_tags:
                return
            if tagl == "table":
                if self._current_table and self._current_table:
                    self.tables.append(self._current_table)
                self._current_table = None
            elif tagl == "tr":
                if self._current_table is not None and self._current_row is not None:
                    self._current_table.append(self._current_row)
                self._current_row = None
            elif tagl in ("td", "th"):
                if self._current_row is not None and self._in_cell:
                    self._current_row.append(
                        (self._current_cell or "").strip()
                    )
                self._in_cell = False
                self._current_cell = None

        def handle_data(self, data):
            if self._in_cell and self._current_cell is not None:
                self._current_cell += data

    extractor = TableExtractor()
    try:
        extractor.feed(html_content)
    except Exception:
        pass

    return extractor.tables if extractor.tables else None


def _html_table_to_pipe(table: list[list[str]]) -> str:
    """将 HTML 提取的二维数组转为 Pandoc 管道表格字符串"""
    if not table:
        return ""
    max_cols = max(len(row) for row in table) if table else 0
    if max_cols < 2:
        return ""

    lines = []
    for row in table:
        cells = []
        for i in range(max_cols):
            cell = row[i] if i < len(row) else ""
            cell = html.unescape(cell)
            cell = re.sub(r"\s+", " ", cell).strip()
            cells.append(cell)
        lines.append("|" + "|".join(cells) + "|")

    sep = "|" + "|".join(["---"] * max_cols) + "|"
    lines.insert(1, sep)
    return "\n".join(lines)


def _html_to_markdown(html_content: str) -> str | None:
    """
    将 HTML 剪贴板内容（含 <table>）转换为 Markdown，
    表格以管道格式内嵌在文档流中，保持原始顺序。

    不依赖 BeautifulSoup，用 Python 标准库 HTMLParser 实现。
    """
    from html.parser import HTMLParser

    # 提取 <!--StartFragment-->...<!--EndFragment--> 片段
    for marker in ("<!--StartFragment-->", "<!--EndFragment-->"):
        pos = html_content.find(marker)
        if pos >= 0:
            if marker == "<!--StartFragment-->":
                html_content = html_content[pos + len(marker):]
            else:  # EndFragment
                html_content = html_content[:pos]

    class HtmlToMarkdown(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts: list[tuple[str, object]] = []
            self._text_buf: list[str] = []
            self._in_table = False
            self._table_buf: list[list[str]] = []
            self._row_buf: list[str] = []
            self._cell_buf: list[str] = []
            self._in_cell = False
            self._skip_depth = 0

        def _flush_text(self):
            text = "".join(self._text_buf)
            text = html.unescape(text)
            text = re.sub(r"[ \t]+", " ", text).strip()
            if text:
                self.parts.append(("text", text))
            self._text_buf = []

        def handle_starttag(self, tag, attrs):
            t = tag.lower()
            if t in ("script", "style", "head", "meta", "link"):
                self._skip_depth += 1
                return
            if self._skip_depth:
                return
            if t == "table":
                self._flush_text()
                self._in_table = True
                self._table_buf = []
                self._row_buf = []
            elif t == "tr" and self._in_table:
                self._row_buf = []
            elif t in ("td", "th") and self._in_table:
                self._in_cell = True
                self._cell_buf = []
            elif t == "br":
                self._text_buf.append("\n")

        def handle_endtag(self, tag):
            t = tag.lower()
            if t in ("script", "style", "head", "meta", "link"):
                self._skip_depth = max(0, self._skip_depth - 1)
                return
            if self._skip_depth:
                return
            if t == "table":
                self._in_table = False
                if self._table_buf and len(self._table_buf) >= 2:
                    self.parts.append(("table", self._table_buf))
                self._table_buf = []
                self._row_buf = []
            elif t == "tr" and self._in_table:
                if self._row_buf:
                    self._table_buf.append(list(self._row_buf))
                self._row_buf = []
            elif t in ("td", "th") and self._in_table and self._in_cell:
                cell_text = "".join(self._cell_buf).strip()
                self._row_buf.append(cell_text)
                self._in_cell = False
                self._cell_buf = []

        def handle_data(self, data):
            if self._skip_depth:
                return
            if self._in_cell:
                self._cell_buf.append(data)
            else:
                self._text_buf.append(data)

    converter = HtmlToMarkdown()
    try:
        converter.feed(html_content)
    except Exception:
        return None
    converter._flush_text()

    # 构建 Markdown 输出
    md_parts = []
    for kind, data in converter.parts:
        if kind == "text":
            md_parts.append(data)
        elif kind == "table" and data:
            max_cols = max(len(r) for r in data) if data else 0
            if max_cols < 2:
                continue
            sep = "|" + "|".join(["---"] * max_cols) + "|"
            pipe_rows = []
            for row in data:
                cells = []
                for i in range(max_cols):
                    cell = row[i] if i < len(row) else ""
                    cell = html.unescape(cell)
                    cell = re.sub(r"\s+", " ", cell).strip()
                    cells.append(cell)
                pipe_rows.append("|" + "|".join(cells) + "|")
            pipe_rows.insert(1, sep)
            md_parts.append("\n".join(pipe_rows))

    result = "\n\n".join(md_parts)
    result = re.sub(r"\n{4,}", "\n\n\n", result)
    return result.strip()


# ============================================================
# 配置
# ============================================================

# 样式配置
STYLES = {
    "heading_font": "仿宋_GB2312",      # 标题中文
    "body_font": "仿宋_GB2312",         # 正文中文
    "western_font": "Times New Roman",  # 英文/数字
    "math_font": "Cambria Math",        # 公式字体（Word OMML 默认）
    "code_font": "Consolas",            # 代码等宽字体
    "code_bg_color": "F5F5F5",
    # 标题字号（仿 AI 显示风格）
    "heading_sizes": {1: 18, 2: 16, 3: 14, 4: 12, 5: 12, 6: 12},
    # 正文
    "body_size": 11,               # 五号（标准正文阅读字号）
    "body_line_spacing": 24,       # 最小值 24pt，含公式时可自动撑开
    "body_line_spacing_rule": "at_least",  # "at_least"=最小值自适应, "exact"=固定值
    "body_space_after": 4,         # 段后间距
    "heading_space_before": 16,    # 标题段前
    "heading_space_after": 8,      # 标题段后
}

# ============================================================
# 中英混排字体工具
# ============================================================


def _set_run_mixed_font(run, east_asia_font=None, western_font=None, size=None):
    """
    设置 Run 的中英混排字体：
      - 中文 → east_asia_font（默认 仿宋_GB2312）
      - 英文/数字 → western_font（默认 Times New Roman）
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ea = east_asia_font or STYLES["body_font"]
    wf = western_font or STYLES["western_font"]

    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)

    rFonts.set(qn("w:ascii"), wf)
    rFonts.set(qn("w:hAnsi"), wf)
    rFonts.set(qn("w:eastAsia"), ea)
    rFonts.set(qn("w:cs"), wf)

    if size is not None:
        run.font.size = size


# ============================================================
# 核心：Pandoc 转换（主引擎）
# ============================================================


def has_pandoc() -> bool:
    """检查 pandoc 是否可用"""
    return shutil.which("pandoc") is not None


def convert_with_pandoc(content: str, output_path: str, debug: bool = False) -> bool:
    """
    使用 Pandoc 将 Markdown+LaTeX 内容转为 Word 文档。

    Pandoc 自动处理：
      - $...$ / $$...$$ LaTeX 数学 → Word OMML 原生公式（可编辑）
      - # ## ### 标题 → Word 标题样式
      - ```code``` 代码块
      - 表格、列表、粗斜体等
    """
    # 预清理：一些 AI 平台的特殊格式
    content = _preprocess_content(content)

    # 表统计
    pipe_tables = content.count("|---")
    if debug or _DEBUG:
        _dump_content_debug(content, output_path, pipe_tables)

    # 用临时文件传递内容（避免 shell 编码问题）
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", encoding="utf-8", delete=False
    ) as f:
        f.write(content)
        input_path = f.name

    try:
        # Pandoc 参数说明：
        #   -f markdown+tex_math_dollars  : 启用 LaTeX 数学公式解析
        #   --wrap=preserve               : 保留原始换行
        cmd = [
            "pandoc",
            input_path,
            "-o", output_path,
            "-f", "markdown+tex_math_dollars+tex_math_single_backslash",
            "-t", "docx",
            "--wrap=preserve",
            "--metadata", "title=AI 内容导出",
        ]

        result = subprocess.run(
            cmd, capture_output=True, text=True, encoding="utf-8"
        )

        if result.returncode != 0:
            print(f"⚠️  Pandoc 转换出错: {result.stderr.strip()}", file=sys.stderr)
            return False

        # 后处理：用 python-docx 微调样式
        _postprocess_docx(output_path)

        return True

    finally:
        # 清理临时文件
        try:
            os.unlink(input_path)
        except OSError:
            pass


def _preprocess_content(content: str) -> str:
    """
    预处理 AI 内容，提升 Pandoc 兼容性。
    """
    content = _fix_chinese_headings(content)
    content = _fix_space_aligned_tables(content)   # 空格/制表符对齐表格 → 管道表格
    content = _fix_markdown_tables(content)         # 补全缺失的分隔行

    lines = content.split("\n")
    processed = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # --- 检测 Mermaid 代码块，替换为普通代码块（Pandoc 会原样保留） ---
        if line.strip().startswith("```mermaid"):
            processed.append("```mermaid")
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                processed.append(lines[i])
                i += 1
            if i < len(lines):
                processed.append(lines[i])  # 结束 ```
            i += 1
            # 添加说明文字
            processed.append("")
            processed.append("> 📊 *上方为 Mermaid 图表源码，可复制到 Mermaid Live Editor 查看*")
            processed.append("")
            continue

        # --- 修复 AI 常见的不规范 Markdown ---
        # 补充标题前的空行（Pandoc 需要）
        if re.match(r"^#{1,6}\s", line):
            if processed and processed[-1].strip() != "":
                processed.append("")

        # --- 修复行内公式中不正确的转义 ---
        # 某些 AI 会输出 \(...\) 而不是 $...$
        # 注意：去除首尾空格，否则 Pandoc 不识别带空格的 $…$
        line = re.sub(r"\\\(\s*(.+?)\s*\\\)", r"$\1$", line)
        line = re.sub(r"\\\[\s*(.+?)\s*\\\]", r"$$\1$$", line)

        processed.append(line)
        i += 1

    return "\n".join(processed)


def _fix_chinese_headings(content: str) -> str:
    """
    智能检测中文纯文本标题（无 # 标记），自动转为 Markdown 标题。

    处理的模式：
      一、二、...        → H2
      第一章/第二节/...  → H2
      依据 N：/方法 N：/... → H3
      方案 X：          → H4

    同时处理合并在同一行的标题（如 "三、xxx方法 1：yyy" 拆成两行）。
    """
    # (正则, 标题级别)
    PATTERNS = [
        (r"^[一二三四五六七八九十]+[、．]\s*", 2),              # 一、二、…
        (r"^第[一二三四五六七八九十\d]+[章节篇部分]\s*", 2),   # 第一章、第二节 …
        (r"^(方法|步骤|阶段|结论|条件|参数)\s*\d+\s*[：:]\s*\S", 3),
        (r"^依据\s+\d+\s*[：:]\s*\S", 3),                       # 依据 2：…（必须有数字）
        (r"^方案\s*[A-D一二三四五六]\s*[：:]\s*\S", 4),         # 方案 A：…
    ]

    # 先拆分合并标题：如 "三、xxx方法 1：yyy" → 在 "方法 1：" 前断行
    # 注意：对已经是 Markdown 标题（以 # 开头）的行不做拆分
    split_patterns = [
        r"(方法\s*\d+\s*[：:])",
        r"(步骤\s*\d+\s*[：:])",
        r"(方案\s*[A-D一二三四五六]\s*[：:])",
        r"(依据\s*\d+\s*[：:])",
    ]
    split_lines = content.split("\n")
    new_lines = []
    for line in split_lines:
        if re.match(r"^#{1,6}\s", line.strip()):
            new_lines.append(line)  # 已有 # 标记，不拆分
            continue
        for pat in split_patterns:
            if re.search(pat, line):
                line = re.sub(pat, r"\n\1", line)
        new_lines.append(line)
    content = "\n".join(new_lines)

    lines = content.split("\n")
    result = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append(line)
            continue
        # 已经是 Markdown 标题就跳过
        if re.match(r"^#{1,6}\s", stripped):
            result.append(line)
            continue

        matched_level = 0
        for pattern, level in PATTERNS:
            if re.match(pattern, stripped):
                matched_level = level
                break

        if matched_level:
            # 确保标题前有空行
            if result and result[-1].strip() != "":
                result.append("")
            result.append(f"{'#' * matched_level} {stripped}")
            result.append("")
        else:
            result.append(line)

    return "\n".join(result)


def _fix_space_aligned_tables(content: str) -> str:
    """
    检测非管道符格式的文本表格，转为 Pandoc 可识别的管道表格（|...|）。

    支持的格式（按优先级）：
      1.  制表符分隔（TSV）—— 每行含至少 1 个 Tab，列数由 \t 拆分决定
      2.  空格对齐 —— 每行含 2+ 连续空格，列数由 [ \t]{{2,}} 拆分决定

    检测条件：
      - 连续 3+ 行，每行能拆出 2+ 列
      - 列数一致（允许 20% 容差）
    """
    import re

    lines = content.split("\n")

    def _table_cols(line: str, method: str = "auto") -> list:
        """按检测方法拆分列"""
        s = line.strip()
        if method == "tab":
            parts = s.split("\t")
        elif method == "space":
            parts = re.split(r"[ \t]{2,}", s)
        else:
            # 自动：优先制表符
            if "\t" in s:
                parts = s.split("\t")
            else:
                parts = re.split(r"[ \t]{2,}", s)
        return [p.strip() for p in parts if p.strip()]

    def _detect_method(line: str) -> str:
        """判断适合的拆分方法"""
        s = line.strip()
        if "\t" in s:
            # 有制表符 → 按 Tab 分
            cols = _table_cols(s, "tab")
            if len(cols) >= 2:
                return "tab"
        if re.search(r"[ \t]{2,}", s):
            cols = _table_cols(s, "space")
            if len(cols) >= 2:
                return "space"
        return ""

    def _is_table_row(line: str) -> bool:
        s = line.strip()
        if not s or s.startswith("|") or s.startswith(">"):
            return False
        # 只要能用任一分隔方式拆出 2+ 列就算候选
        cols = _table_cols(s, "auto")
        return len(cols) >= 2

    result = []
    i = 0
    while i < len(lines):
        if _is_table_row(lines[i]):
            # 确定本块的拆分方法（取第一行的）
            method = _detect_method(lines[i])
            if not method:
                result.append(lines[i])
                i += 1
                continue

            block = [lines[i]]
            i += 1
            while i < len(lines) and _detect_method(lines[i]) == method:
                block.append(lines[i])
                i += 1

            if len(block) >= 3:
                col_counts = [len(_table_cols(l, method)) for l in block]
                most_common = max(set(col_counts), key=col_counts.count)
                consistent = sum(1 for c in col_counts if c == most_common)
                if most_common >= 2 and consistent >= len(block) * 0.8:
                    sep = "|" + "|".join(["---"] * most_common) + "|"
                    pipe_rows = []
                    for row in block:
                        cols = _table_cols(row, method)
                        while len(cols) < most_common:
                            cols.append("")
                        pipe_rows.append("|" + "|".join(cols[:most_common]) + "|")
                    pipe_rows.insert(1, sep)

                    if result and result[-1].strip() != "":
                        result.append("")
                    for pr in pipe_rows:
                        result.append(pr)
                    result.append("")
                    continue

            # 不满足表格条件，原样输出
            for row in block:
                result.append(row)
        else:
            result.append(lines[i])
            i += 1

    return "\n".join(result)


def _fix_markdown_tables(content: str) -> str:
    """
    修复 AI 输出的不完整 Markdown 表格。

    问题：有些 AI 输出表格时省略了分隔行（|---|---|），
    Pandoc 无法识别为表格。

    修复：检测连续以 | 开头的行，如果缺少 --- 行则自动添加。
    """
    lines = content.split("\n")
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 检测是否为表格行（以 | 开头且包含至少 3 个 |）
        if stripped.startswith("|") and stripped.count("|") >= 3:
            table_rows = []
            while i < len(lines):
                ls = lines[i].strip()
                if ls.startswith("|") and ls.count("|") >= 3:
                    table_rows.append(ls)
                    i += 1
                else:
                    break

            # 检查是否有分隔行（包含 --- 内容）
            has_separator = any("---" in row for row in table_rows)

            if not has_separator and len(table_rows) >= 2:
                # 按 | 拆分后取非空段数即为列数（兼容有无末尾 | 两种写法）
                cols = len([c for c in table_rows[0].split("|") if c.strip()])
                separator = "|" + "|".join(["---"] * cols) + "|"
                table_rows.insert(1, separator)

            # 确保表格前后有空行
            if result and result[-1].strip() != "":
                result.append("")
            for row in table_rows:
                result.append(row)
            result.append("")
        else:
            result.append(line)
            i += 1

    return "\n".join(result)


# ============================================================
# 后处理：python-docx 微调
# ============================================================


def _postprocess_docx(docx_path: str):
    """
    用 python-docx 对 Pandoc 输出的文档进行样式微调。
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return  # 没有 python-docx 就跳过后处理

    doc = Document(docx_path)

    # 遍历所有段落，调整样式
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        # --- 调整标题字体 ---
        if style_name.startswith("Heading"):
            level = style_name.replace("Heading", "").strip()
            level = int(level) if level.isdigit() else 1
            _set_heading_style(para, level)
            continue

        # --- 调整代码块样式（Pandoc 通常用 "Source Code" 样式） ---
        if "Code" in style_name or "代码" in style_name:
            _set_code_style(para)
            continue

        # --- 正文样式（Pandoc 可能用 Normal / First Paragraph / Body Text） ---
        if style_name in ("Normal", "First Paragraph", "Body Text") or not style_name:
            _set_body_style(para)

    # --- 调整表格样式 ---
    for table in doc.tables:
        _style_table(table)

    # --- 设置页边距（国标论文格式） ---
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    doc.save(docx_path)


def _set_heading_style(para, level: int):
    """设置标题样式（论文格式：黑色、仿宋_GB2312 + Times New Roman 混排）"""
    from docx.shared import Pt

    size = STYLES["heading_sizes"].get(level, 12)

    for run in para.runs:
        run.font.bold = True
        run.font.color.rgb = None  # 清除颜色，使用默认黑色
        # 中英混排
        _set_run_mixed_font(run, STYLES["heading_font"], STYLES["western_font"], Pt(size))

    para.paragraph_format.space_before = Pt(STYLES["heading_space_before"])
    para.paragraph_format.space_after = Pt(STYLES["heading_space_after"])


def _set_body_style(para):
    """设置正文字体（仿宋_GB2312 + Times New Roman 混排）"""
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    for run in para.runs:
        _set_run_mixed_font(run, STYLES["body_font"], STYLES["western_font"], Pt(STYLES["body_size"]))

    # 行距：使用"至少"模式，普通文本保持 24pt，含分式等高等公式时自动撑开
    rule = STYLES.get("body_line_spacing_rule", "at_least")
    if rule == "at_least":
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    else:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(STYLES["body_line_spacing"])
    para.paragraph_format.space_after = Pt(STYLES["body_space_after"])


def _set_code_style(para):
    """设置代码块样式（黑色，Consolas）"""
    from docx.shared import Pt

    for run in para.runs:
        run.font.name = STYLES["code_font"]
        run.font.size = Pt(9.5)
        run.font.color.rgb = None  # 黑色

    para.paragraph_format.line_spacing = Pt(16)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)


def _style_table(table):
    """美化表格"""
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # 边框
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "4472C4")
        borders.append(element)
    tblPr.append(borders)

    # 表头样式
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "4472C4")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)


# ============================================================
# 纯 Python 降级方案（无 Pandoc 时使用）
# ============================================================


def convert_with_python(content: str, output_path: str):
    """
    不使用 Pandoc，纯 Python + python-docx 构建 Word 文档。

    优点：无需安装 Pandoc
    缺点：LaTeX 公式保留为源代码文本（不会转为可编辑公式）
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ 需要 python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # 设置默认中英混排字体（在 Normal 样式上）
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style = doc.styles["Normal"]
    style.font.size = Pt(STYLES["body_size"])
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), STYLES["western_font"])
    rFonts.set(qn("w:hAnsi"), STYLES["western_font"])
    rFonts.set(qn("w:eastAsia"), STYLES["body_font"])
    rFonts.set(qn("w:cs"), STYLES["western_font"])
    style.paragraph_format.line_spacing = Pt(STYLES["body_line_spacing"])
    style.paragraph_format.space_after = Pt(STYLES["body_space_after"])

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # --- 代码块处理 ---
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip().strip("`").strip()
                code_buffer = []
            else:
                in_code_block = False
                _add_code_block(doc, "\n".join(code_buffer), code_lang)
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # --- 空行 ---
        if line.strip() == "":
            i += 1
            continue

        # --- 标题 ---
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # 去掉行内 Markdown 标记（粗体等）
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            heading = doc.add_heading(text, level=level if level <= 4 else 4)
            _set_heading_style(heading, level)
            i += 1
            continue

        # --- 水平线 ---
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("").paragraph_format.space_before = Pt(6)
            # 添加一条水平线（使用下划线段落）
            p = doc.add_paragraph("")
            run = p.add_run("_" * 80)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(6)
            i += 1
            continue

        # --- 引用块 ---
        if line.strip().startswith(">"):
            quote_text = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_text.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            _add_quote_block(doc, "\n".join(quote_text))
            continue

        # --- 列表（无序） ---
        if re.match(r"^[\s]*[-*+]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[\s]*[-*+]\s+", lines[i]):
                items.append(re.sub(r"^[\s]*[-*+]\s+", "", lines[i]))
                i += 1
            for item in items:
                p = doc.add_paragraph(item, style="List Bullet")
                _set_body_style(p)
            continue

        # --- 列表（有序） ---
        if re.match(r"^[\s]*\d+[\.\)]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^[\s]*\d+[\.\)]\s+", lines[i]):
                items.append(re.sub(r"^[\s]*\d+[\.\)]\s+", "", lines[i]))
                i += 1
            for idx, item in enumerate(items):
                p = doc.add_paragraph(item, style="List Number")
                _set_body_style(p)
            continue

        # --- 普通段落（含行内格式处理） ---
        para = doc.add_paragraph()
        _process_inline_formatting(para, line)
        _set_body_style(para)
        i += 1

    # 保存文档
    doc.save(output_path)


def _process_inline_formatting(para, text: str):
    """
    处理行内格式：粗体、斜体、行内代码、行内公式。
    """
    from docx.shared import Pt, RGBColor

    # 用正则分段：将文本按特殊标记拆分成多个片段
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\$(.+?)\$|\\\((.+?)\\\))"
    parts = re.split(pattern, text)

    i = 0
    while i < len(parts):
        segment = parts[i]
        if segment is None or segment == "":
            i += 1
            continue

        # 判断这个片段是哪一类的
        if i + 1 < len(parts) and parts[i] is not None:
            # 检查是否匹配到模式
            full_match = parts[i]
            # 粗体
            bold_match = re.match(r"^\*\*(.+)\*\*$", segment)
            if bold_match:
                run = para.add_run(bold_match.group(1))
                run.bold = True
                i += 1
                continue

            # 斜体
            italic_match = re.match(r"^\*(.+)\*$", segment)
            if italic_match:
                run = para.add_run(italic_match.group(1))
                run.italic = True
                i += 1
                continue

            # 行内代码
            code_match = re.match(r"^`(.+)`$", segment)
            if code_match:
                run = para.add_run(code_match.group(1))
                run.font.name = STYLES["code_font"]
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                i += 1
                continue

            # 公式 $...$ （降级：保留源码，用特殊颜色标注）
            math_match = re.match(r"^\$(.+)\$$", segment)
            if math_match:
                run = para.add_run(f" {math_match.group(1)} ")
                run.font.name = STYLES["code_font"]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                run.italic = True
                i += 1
                continue

        # 普通文本
        run = para.add_run(segment)
        _set_run_mixed_font(run, STYLES["body_font"], STYLES["western_font"], Pt(STYLES["body_size"]))

        i += 1


def _add_code_block(doc, code: str, lang: str):
    """添加代码块段落"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if lang:
        # 添加语言标签
        label = doc.add_paragraph()
        run = label.add_run(f"  [{lang}]")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = STYLES["code_font"]
        label.paragraph_format.space_after = Pt(0)
        label.paragraph_format.space_before = Pt(4)

    for code_line in code.split("\n"):
        para = doc.add_paragraph()
        run = para.add_run(code_line)
        run.font.name = STYLES["code_font"]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # 设置缩进
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(15)

        # 添加灰色背景
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        shading.set(qn("w:val"), "clear")
        para.paragraph_format.element.get_or_add_pPr().append(shading)

    # 代码块后间距
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def _add_quote_block(doc, text: str):
    """添加引用块"""
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    # 左侧灰色竖线（通过段落的左边框模拟引用效果）
    pPr = para.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "4472C4")
    pBdr.append(left)
    pPr.append(pBdr)

    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)


# ============================================================
# 辅助功能
# ============================================================


def _dump_content_debug(content: str, output_path: str, table_count: int):
    """调试模式：保存预处理后的 markdown 并输出诊断信息"""
    from pathlib import Path

    out = Path(output_path)
    debug_md = out.with_suffix(".debug.md")

    # 写调试文件
    debug_md.write_text(content, encoding="utf-8")
    print(f"\n🐛 [DEBUG] 预处理后的 Markdown 已保存到: {debug_md}")

    # 表格统计
    if table_count:
        # 找每张表
        table_starts = []
        for i, line in enumerate(content.split("\n")):
            if line.lstrip().startswith("|") and "---" not in line and not line.lstrip().startswith("|---"):
                # 检查上一行不是管道行 → 新表格的第一行
                prev = content.split("\n")[i-1] if i > 0 else ""
                if not prev.lstrip().startswith("|"):
                    table_starts.append(i)
        print(f"   📊 管道表格: {table_count} 个分隔行 → ~{len(table_starts)} 张表")
    else:
        print("   ⚠️  未检测到任何管道表格！")

    # 检查原始内容中是否有 Tab
    tab_lines = sum(1 for l in content.split("\n") if "\t" in l)
    print(f"   🔤 含制表符行数: {tab_lines}")

    # 预览前 10 行管道表格行
    pipe_lines = [(i, l) for i, l in enumerate(content.split("\n")) if l.startswith("|")]
    if pipe_lines:
        print(f"   📝 管道表格行（前 10 行）:")
        for idx, line in pipe_lines[:10]:
            print(f"      L{idx}: {line[:120]}")
    else:
        print("   📝 未找到管道表格行")
    print()


def detect_mermaid(content: str) -> bool:
    """检测内容中是否包含 Mermaid 图表"""
    return "```mermaid" in content


def detect_math(content: str) -> dict:
    """检测数学公式内容"""
    inline = len(re.findall(r"(?<!\$)\$[^$]+\$(?!\$)", content))
    display = len(re.findall(r"\$\$[^$]+\$\$", content))
    paren = len(re.findall(r"\\\(.*?\\\)", content))
    bracket = len(re.findall(r"\\\[.*?\\\]", content))
    return {
        "inline": inline,
        "display": display,
        "paren": paren,
        "bracket": bracket,
        "total": inline + display + paren + bracket,
    }


def format_detection_report(content: str) -> list:
    """生成内容检测报告"""
    reports = []

    lines = content.count("\n") + 1
    chars = len(content)
    reports.append(f"📄 总行数: {lines} | 总字符: {chars}")

    math = detect_math(content)
    if math["total"] > 0:
        reports.append(
            f"📐 LaTeX 公式: {math['total']} 处"
            f"（内联 {math['inline']} | 展示 {math['display']} | \\(\\) {math['paren']} | \\[\\] {math['bracket']}）"
        )

    # 检测代码块
    code_blocks = re.findall(r"```(\w*)", content)
    if code_blocks:
        lang_counts = {}
        for lang in code_blocks:
            l = lang if lang else "未指定"
            lang_counts[l] = lang_counts.get(l, 0) + 1
        reports.append(
            f"💻 代码块: {len(code_blocks)} 个"
            + (f"（{', '.join(f'{k}×{v}' for k, v in lang_counts.items())}）" if lang_counts else "")
        )

    # 检测标题
    headings = re.findall(r"^(#{1,6})\s", content, re.MULTILINE)
    if headings:
        h_levels = {}
        for h in headings:
            h_levels[h] = h_levels.get(h, 0) + 1
        reports.append(
            f"📑 标题: {len(headings)} 个"
            + (f"（{', '.join(f'H{len(k)}×{v}' for k, v in sorted(h_levels.items()))}）" if h_levels else "")
        )

    if detect_mermaid(content):
        reports.append("📊 Mermaid 图表: 已检测到（源码将保留在文档中）")

    return reports


# ============================================================
# 主入口
# ============================================================


def main():
    global _DEBUG

    parser = argparse.ArgumentParser(
        description="AI2DOCX — AI 内容一键转为 Word 文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python ai2docx.py                       从剪贴板读取 → 输出 output.docx
  python ai2docx.py conversation.md       从 Markdown 文件读取
  python ai2docx.py input.md output.docx  指定输入输出
  python ai2docx.py --out 论文.docx       从剪贴板读取 → 输出到 论文.docx
  python ai2docx.py --debug               调试模式：保存中间文件 + 诊断输出
        """,
    )
    parser.add_argument("input", nargs="?", help="输入 Markdown 文件（省略则从剪贴板读取）")
    parser.add_argument("output", nargs="?", default=None, help="输出 .docx 文件路径")
    parser.add_argument(
        "--out", "-o", dest="output_opt", default=None, help="输出文件路径（从剪贴板读取时使用）"
    )
    parser.add_argument(
        "--quiet", "-q", action="store_true", help="静默模式（不输出检测报告）"
    )
    parser.add_argument(
        "--debug", "-d", action="store_true", help="调试模式：保存预处理后的 .debug.md 并输出诊断"
    )
    args = parser.parse_args()

    _DEBUG = args.debug

    # ---------- 读取输入 ----------
    if args.input:
        input_path = Path(args.input)
        if not input_path.exists():
            print(f"❌ 文件不存在: {args.input}", file=sys.stderr)
            sys.exit(1)
        content = input_path.read_text(encoding="utf-8")
        source_desc = f"文件 {args.input}"
    else:
        try:
            import pyperclip
            content = pyperclip.paste()
        except ImportError:
            print("❌ 需要 pyperclip: pip install pyperclip", file=sys.stderr)
            sys.exit(1)

        if not content or not content.strip():
            print("❌ 剪贴板为空！请先复制 AI 内容", file=sys.stderr)
            sys.exit(1)
        source_desc = "剪贴板"

        # ── 调试：查看剪贴板原始格式 ──
        if _DEBUG:
            print(f"\n🐛 [DEBUG] 剪贴板诊断")
            print(f"   总字符: {len(content)}")
            # 显示不可见字符
            esc = content[:500].replace("\t", "␣[TAB]␣").replace("\n", "␣[LF]␣\n")
            print(f"   前 500 字符（转义后）:\n{esc}")
            html_raw = _get_html_from_clipboard()
            if html_raw:
                print(f"   HTML 剪贴板: {len(html_raw)} 字符")
                from html.parser import HTMLParser
                tables = _extract_tables_from_html(html_raw)
                if tables:
                    print(f"   HTML 中提取到 {len(tables)} 张表:")
                    for ti, t in enumerate(tables):
                        print(f"      表{ti+1}: {len(t)} 行 × {max(len(r) for r in t) if t else 0} 列")
                else:
                    print(f"   HTML 中未提取到表格")
            else:
                print(f"   HTML 剪贴板: 不可用")
            print()

        # ── HTML 剪贴板回退：纯文本无表格结构时，从 HTML 生成完整 Markdown ──
        pipe_count_check = content.count("|---")
        if pipe_count_check == 0:
            html_content = _get_html_from_clipboard()
            if html_content:
                md_from_html = _html_to_markdown(html_content)
                if md_from_html and len(md_from_html) > len(content) * 0.3:
                    print(f"🔍 纯文本无表格结构，从 HTML 剪贴板生成 Markdown...")
                    content = md_from_html
                    print(f"   ✅ 已转换（含内联表格，共 {content.count('|---')} 张表）")
                else:
                    # HTML 转换结果太短，不安全；回退用纯文本 + 追加表格
                    tables = _extract_tables_from_html(html_content)
                    if tables:
                        print(f"🔍 纯文本无表格，从 HTML 提取 {len(tables)} 张表追加到末尾...")
                        pipe_list = [_html_table_to_pipe(t) for t in tables if _html_table_to_pipe(t)]
                        if pipe_list:
                            content = content.rstrip() + "\n\n" + "\n\n".join(pipe_list) + "\n"
                            print(f"   ⚠️  追加模式（表格在文末）{len(pipe_list)} 张表")

    # ---------- 内容检测报告 ----------
    if not args.quiet:
        print(f"\n{'='*50}")
        print(f"📥 已从 {source_desc} 读取内容")
        print(f"{'='*50}")
        for report in format_detection_report(content):
            print(f"   {report}")
        # 管道表格统计
        pipe_count = content.count("|---")
        if pipe_count:
            print(f"   📊 管道表格: {pipe_count} 张")
        print()

    # ---------- 确定输出路径 ----------
    if args.output:
        output_path = args.output
    elif args.output_opt:
        output_path = args.output_opt
    else:
        output_path = "output.docx"

    output_path = Path(output_path)
    if output_path.suffix.lower() not in (".docx",):
        output_path = output_path.with_suffix(".docx")

    # 避免覆盖已有文件
    if output_path.exists():
        base = output_path.stem
        parent = output_path.parent
        counter = 1
        while output_path.exists():
            output_path = parent / f"{base}_{counter}.docx"
            counter += 1
        print(f"⚠️  文件已存在，自动重命名为: {output_path.name}")

    # ---------- 执行转换 ----------
    if has_pandoc():
        if not args.quiet:
            print("🔧 使用 Pandoc 引擎（支持 LaTeX → Word 原生公式）...")
        success = convert_with_pandoc(content, str(output_path), debug=args.debug)
        if not success:
            print("⚠️  Pandoc 转换失败，尝试降级方案...", file=sys.stderr)
            convert_with_python(content, str(output_path))
    else:
        if not args.quiet:
            print("🔧 Pandoc 未安装，使用 Python 降级引擎...")
            print("   ⚠  建议安装 Pandoc 以获得 LaTeX 公式支持: https://pandoc.org")
        convert_with_python(content, str(output_path))

    # ---------- 完成 ----------
    size = output_path.stat().st_size
    size_str = f"{size / 1024:.1f} KB" if size < 1024 * 1024 else f"{size / 1024 / 1024:.1f} MB"

    print(f"\n{'='*50}")
    print(f"✅ 转换完成！")
    print(f"   📄 输出: {output_path.resolve()}")
    print(f"   📏 大小: {size_str}")
    print(f"{'='*50}")

    # 如果检测到 Mermaid，给出提示
    if detect_mermaid(content):
        print()
        print("💡 提示: 文档中包含 Mermaid 图表源码。")
        print("   可访问 https://mermaid.live 粘贴源码查看渲染效果。")

    # 如果检测到公式但使用的是 Python 降级，给出提示
    math = detect_math(content)
    if math["total"] > 0 and not has_pandoc():
        print()
        print("💡 提示: 文档中包含 LaTeX 数学公式源码。")
        print("   安装 Pandoc (https://pandoc.org) 后运行可自动转为 Word 原生公式。")

    # 调试提示
    if _DEBUG:
        debug_md = output_path.with_suffix(".debug.md")
        if debug_md.exists():
            print(f"\n🐛 调试文件: {debug_md}")


if __name__ == "__main__":
    main()
